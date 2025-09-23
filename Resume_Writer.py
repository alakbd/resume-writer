#!/usr/bin/env python
# coding: utf-8

# In[1]:


"""
Enhanced Resume Writer with Professional DOCX & PDF Export
Includes sample preview feature for demonstration
"""

import streamlit as st
import openai
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
import tempfile
import re
import time
from typing import Tuple, Optional

# Import for PDF text extraction
try:
    import PyPDF2
    from io import BytesIO
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False
    st.warning("PDF support requires PyPDF2. Install with: pip install PyPDF2")

# -------- Sample Data --------
SAMPLE_CV = """
John Doe
Email: john.doe@example.com | Phone: (555) 123-4567 | LinkedIn: linkedin.com/in/johndoe

PROFESSIONAL SUMMARY
Software Engineer with 5+ years of experience in full-stack development. Specialized in JavaScript, Python, and cloud technologies. Proven track record of delivering scalable web applications and improving system performance.

EXPERIENCE
Senior Software Engineer, Tech Solutions Inc. (2020-Present)
- Developed and maintained responsive web applications using React.js and Node.js
- Led a team of 4 developers to deliver a customer portal that increased user engagement by 35%
- Implemented CI/CD pipelines reducing deployment time by 40%
- Optimized database queries, improving application performance by 25%

Software Developer, Innovate Labs (2018-2020)
- Built RESTful APIs using Python and Django framework
- Collaborated with UX designers to implement user-friendly interfaces
- Reduced server costs by 20% through code optimization and caching strategies

EDUCATION
Bachelor of Science in Computer Science
University of Technology, Graduated 2018

SKILLS
- Programming Languages: JavaScript, Python, Java, SQL
- Frameworks: React, Node.js, Django, Express.js
- Tools: Git, Docker, AWS, Jenkins, MongoDB
- Certifications: AWS Certified Developer, Scrum Master
"""

SAMPLE_JOB_DESCRIPTION = """
Senior Full Stack Developer Position

We are looking for an experienced Senior Full Stack Developer to join our dynamic team. The ideal candidate will have a strong background in both front-end and back-end development, with expertise in modern JavaScript frameworks and cloud technologies.

Responsibilities:
- Design, develop, and maintain scalable web applications
- Collaborate with cross-functional teams to define, design, and ship new features
- Write clean, maintainable, and efficient code
- Implement security and data protection measures
- Optimize applications for maximum speed and scalability
- Mentor junior developers and conduct code reviews

Requirements:
- 5+ years of experience in full-stack development
- Proficiency with JavaScript frameworks (React, Angular, or Vue)
- Strong experience with Node.js and Python
- Familiarity with database technology such as MongoDB and PostgreSQL
- Experience with cloud services (AWS, Azure, or GCP)
- Knowledge of code versioning tools, such as Git
- Excellent problem-solving skills and attention to detail

Preferred Qualifications:
- Experience with containerization technologies (Docker, Kubernetes)
- Familiarity with CI/CD pipelines
- Understanding of agile methodologies
- Bachelor's degree in Computer Science or related field

Benefits:
- Competitive salary and performance bonuses
- Comprehensive health insurance
- Flexible working hours and remote work options
- Professional development opportunities
"""

# -------- Configuration --------
DEFAULT_FONT = "Helvetica"
DEFAULT_FONT_SIZE = 11
HEADING_FONT_SIZE = 14
LINE_SPACING = 1.2

# -------- Enhanced Prompt Engineering --------
def build_prompt(resume_text: str, job_text: str, tone: str = "Professional") -> str:
    """Enhanced prompt for ATS-optimized, job-aligned résumé."""
    prompt = f"""
You are an expert professional résumé writer with deep knowledge of Applicant Tracking Systems (ATS). 
Rewrite the candidate's résumé to maximize their chances for the specific job while maintaining truthfulness.

CRITICAL GUIDELINES:
1. Job Alignment: 
   - Extract key skills, technologies, and requirements from the job description
   - Mirror the language and terminology used in the job description
   - Prioritize relevant experience and quantify achievements with metrics where possible

2. ATS Optimization:
   - Include relevant keywords from the job description naturally
   - Use standard section headers (Professional Summary, Experience, Education, Skills, Certifications)
   - Use bullet points for achievements with action verbs (Managed, Developed, Increased, Reduced)
   - Avoid tables, columns, graphics, or unusual formatting

3. Content Structure:
   - Professional Summary: 3-4 lines highlighting most relevant qualifications
   - Experience: Focus on last 10-15 years, emphasize relevant roles
   - Skills: Categorize (Technical, Soft, Certifications) and match job requirements
   - Keep to 1-2 pages maximum

4. Tone: Write in a {tone.lower()} tone.

Candidate's current résumé:
{resume_text}

Job description:
{job_text}

Generate ONLY the rewritten résumé with no explanations or commentary:
"""
    return prompt

# -------- Robust OpenAI API Call --------
def call_openai_chat(prompt: str, api_key: str, max_retries: int = 3) -> str:
    """Enhanced OpenAI API call with retry logic and better error handling."""
    openai.api_key = api_key
    for attempt in range(max_retries):
        try:
            response = openai.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": "You are a professional resume writer specializing in ATS optimization."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.3,
                max_tokens=2500
            )
            return response.choices[0].message.content.strip()
        except openai.error.RateLimitError:
            if attempt < max_retries - 1:
                wait_time = 2 ** attempt  # Exponential backoff
                st.warning(f"Rate limit exceeded. Retrying in {wait_time} seconds...")
                time.sleep(wait_time)
            else:
                return "Error: OpenAI API rate limit exceeded. Please try again later."
        except openai.error.AuthenticationError:
            return "Error: Invalid API key. Please check your OpenAI API credentials."
        except openai.error.InvalidRequestError as e:
            return f"Error: Invalid request to OpenAI API: {str(e)}"
        except Exception as e:
            if attempt < max_retries - 1:
                time.sleep(1)
            else:
                return f"Error: Failed to generate resume after {max_retries} attempts. {str(e)}"
    return "Error: Unexpected error occurred during resume generation."

# -------- Always return "Curricula Vitae" --------
def extract_name(resume_text: str) -> str:
    """Return 'Curricula Vitae' as the title for all resumes."""
    return "Curricula Vitae"

# -------- PDF Text Extraction --------
def extract_text_from_pdf(file) -> str:
    """Extract text content from a PDF file."""
    if not PDF_SUPPORT:
        return "Error: PDF support not available. Please install PyPDF2."
    
    try:
        pdf_reader = PyPDF2.PdfReader(file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        return f"Error reading PDF: {str(e)}"

# -------- Enhanced Word Document Generation --------
def save_resume_docx(resume_text: str, filename: str = "resume.docx") -> str:
    """Create a professionally formatted Word document with improved styling."""
    doc = Document()
    
    # Set default font for the document
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(DEFAULT_FONT_SIZE)
    
    # Add title
    candidate_name = extract_name(resume_text)
    title = doc.add_paragraph(candidate_name)
    title.style = doc.styles['Title']
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph()  # Add spacing
    
    # Process resume content
    lines = resume_text.split("\n")
    current_section = None
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Detect section headers
        if re.match(r'^(PROFESSIONAL SUMMARY|EXPERIENCE|EDUCATION|SKILLS|CERTIFICATIONS|PROJECTS)$', line, re.IGNORECASE):
            current_section = line.upper()
            
            # Add section header
            para = doc.add_paragraph()
            run = para.add_run(current_section)
            run.bold = True
            run.font.size = Pt(HEADING_FONT_SIZE)
            run.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
            para.paragraph_format.space_after = Pt(6)
            para.paragraph_format.space_before = Pt(12)
            
        # Process bullet points (especially in experience section)
        elif current_section == "EXPERIENCE" and (line.startswith('-') or line.startswith('•') or line.startswith('*')):
            bullet_text = re.sub(r'^[-•*]\s*', '', line)  # Remove bullet characters
            para = doc.add_paragraph(bullet_text, style='List Bullet')
            para.paragraph_format.space_after = Pt(3)
            para.paragraph_format.left_indent = Pt(18)
            
        # Process regular content
        else:
            para = doc.add_paragraph(line)
            para.paragraph_format.space_after = Pt(3)
    
    doc.save(filename)
    return filename

# -------- Enhanced PDF Generation --------
def save_resume_pdf(resume_text: str, filename: str = "resume.pdf") -> str:
    """Create a professionally formatted PDF with improved styling."""
    doc = SimpleDocTemplate(
        filename, 
        pagesize=A4,
        rightMargin=50, 
        leftMargin=50, 
        topMargin=50, 
        bottomMargin=50
    )
    
    # Create custom styles without conflicting with existing ones
    styles = getSampleStyleSheet()
    
    # Create unique style names to avoid conflicts
    resume_styles = {
        'Body': ParagraphStyle(
            name='ResumeBody',
            parent=styles['Normal'],
            fontName=DEFAULT_FONT,
            fontSize=DEFAULT_FONT_SIZE,
            leading=DEFAULT_FONT_SIZE * LINE_SPACING,
            spaceAfter=6
        ),
        'Heading': ParagraphStyle(
            name='ResumeHeading',
            parent=styles['Heading2'],
            fontName=f'{DEFAULT_FONT}-Bold',
            fontSize=HEADING_FONT_SIZE,
            textColor='#003366',  # Dark blue
            spaceAfter=12,
            spaceBefore=18
        ),
        'Bullet': ParagraphStyle(
            name='ResumeBullet',
            parent=styles['Normal'],
            leftIndent=18,
            bulletIndent=0,
            spaceAfter=3
        ),
        'Title': ParagraphStyle(
            name='ResumeTitle',
            parent=styles['Heading1'],
            fontName=f'{DEFAULT_FONT}-Bold',
            fontSize=18,
            alignment=TA_CENTER,
            spaceAfter=12
        )
    }
    
    # Extract candidate name
    candidate_name = extract_name(resume_text)
    
    # Build story
    story = []
    
    # Add title
    story.append(Paragraph(candidate_name, resume_styles['Title']))
    
    # Process content
    lines = resume_text.split("\n")
    current_section = None
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Detect section headers
        if re.match(r'^(PROFESSIONAL SUMMARY|EXPERIENCE|EDUCATION|SKILLS|CERTIFICATIONS|PROJECTS)$', line, re.IGNORECASE):
            current_section = line.upper()
            story.append(Paragraph(current_section, resume_styles['Heading']))
            
        # Process bullet points
        elif current_section == "EXPERIENCE" and (line.startswith('-') or line.startswith('•') or line.startswith('*')):
            bullet_text = re.sub(r'^[-•*]\s*', '', line)
            story.append(Paragraph(f"• {bullet_text}", resume_styles['Bullet']))
            
        # Process regular content
        else:
            story.append(Paragraph(line, resume_styles['Body']))
    
    doc.build(story)
    return filename

# -------- Sample Preview Function --------
def show_sample_preview():
    """Display a sample preview of how the app works."""
    st.subheader("🎯 Sample Preview")
    st.info("See how the app works with our sample data. No uploads required!")
    
    # Show sample data in expandable sections
    with st.expander("View Sample CV"):
        st.text(SAMPLE_CV)
    
    with st.expander("View Sample Job Description"):
        st.text(SAMPLE_JOB_DESCRIPTION)
    
    # Add a button to generate sample output
    if st.button("🔄 Generate Sample Resume", key="sample_generate"):
        with st.spinner("Generating sample resume..."):
            # Build prompt and call OpenAI
            prompt = build_prompt(SAMPLE_CV, SAMPLE_JOB_DESCRIPTION, tone="Professional")
            api_key_input = os.getenv("OPENAI_API_KEY")
            
            if not api_key_input:
                st.error("OpenAI API key not configured. Sample generation requires a valid API key.")
                return
                
            output = call_openai_chat(prompt, api_key_input)
            
            # Check for errors in API response
            if output.startswith("Error:"):
                st.error(f"Sample generation failed: {output}")
                return
            
            # Display generated resume
            st.subheader("📋 Sample Generated Résumé")
            st.text_area("", output, height=400, label_visibility="collapsed", key="sample_output")
            
            # Create download buttons for sample
            with tempfile.TemporaryDirectory() as tmpdir:
                try:
                    docx_file = save_resume_docx(output, f"{tmpdir}/sample_resume.docx")
                    with open(docx_file, "rb") as f:
                        st.download_button(
                            "📝 Download Sample Word Document", 
                            f, 
                            file_name="sample_resume.docx",
                            help="Download sample in Microsoft Word format",
                            key="sample_docx"
                        )
                except Exception as e:
                    st.error(f"Error creating sample Word document: {str(e)}")
                
                try:
                    pdf_file = save_resume_pdf(output, f"{tmpdir}/sample_resume.pdf")
                    with open(pdf_file, "rb") as f:
                        st.download_button(
                            "📄 Download Sample PDF", 
                            f, 
                            file_name="sample_resume.pdf",
                            help="Download sample in PDF format",
                            key="sample_pdf"
                        )
                except Exception as e:
                    st.error(f"Error creating sample PDF document: {str(e)}")

# -------- Improved Streamlit UI --------
def main():
    st.set_page_config(
        page_title="Professional Résumé Writer", 
        page_icon="📄", 
        layout="centered",
        initial_sidebar_state="expanded"
    )
    
    # Sidebar for instructions and information
    with st.sidebar:
        st.title("ℹ️ Instructions")
        st.info("""
        1. Upload your current résumé (TXT, DOCX, or PDF)
        2. Upload the job description (TXT, DOCX, or PDF)
        3. Select your preferred tone
        4. Click 'Generate Tailored Résumé'
        5. Download your enhanced résumé in Word or PDF format
        """)
        
        if not PDF_SUPPORT:
            st.warning("PDF file support requires PyPDF2. Install with: pip install PyPDF2")
        
        st.title("🔒 Privacy Notice")
        st.caption("""
        - Your documents are processed securely
        - No data is stored on our servers after processing
        - Always review generated content for accuracy
        """)
    
    # Main content area
    st.title("📄 Professional Résumé Writer")
    st.markdown("Transform your résumé into an **ATS-optimized**, job-tailored document that gets noticed.")
    
    # Show sample preview section
    show_sample_preview()
    
    st.markdown("---")
    st.subheader("🚀 Create Your Own Tailored Résumé")
    
    # API key handling
    api_key_input = os.getenv("OPENAI_API_KEY")
    if not api_key_input:
        st.error("OpenAI API key not configured. Please set the OPENAI_API_KEY environment variable.")
        st.stop()
    
    # File upload with better UX
    col1, col2 = st.columns(2)
    
    with col1:
        resume_file = st.file_uploader(
            "Upload Your Résumé", 
            type=["txt", "docx", "pdf"],
            help="Upload your current résumé in TXT, DOCX, or PDF format"
        )
        
    with col2:
        job_file = st.file_uploader(
            "Upload Job Description", 
            type=["txt", "docx", "pdf"],
            help="Upload the job description you're applying for"
        )
    
    # Additional options
    tone = st.selectbox(
        "Select Tone", 
        ["Professional", "Concise", "Achievement-Oriented", "Leadership-Focused"],
        help="Choose the writing style for your résumé"
    )
    
    # Add a preview option
    show_preview = st.checkbox("Show formatted preview before downloading")
    
    # File processing function
    def read_file(file) -> str:
        if file is None:
            return ""
        if file.name.endswith(".txt"):
            return file.read().decode("utf-8")
        elif file.name.endswith(".docx"):
            try:
                doc = Document(file)
                return "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            except Exception as e:
                st.error(f"Error reading DOCX file: {str(e)}")
                return ""
        elif file.name.endswith(".pdf"):
            if not PDF_SUPPORT:
                return "Error: PDF support not available. Please install PyPDF2."
            # Reset file pointer to beginning
            file.seek(0)
            return extract_text_from_pdf(file)
        return ""
    
    # Generate button with improved feedback
    if st.button("✨ Generate Tailored Résumé", type="primary", use_container_width=True):
        if not resume_file or not job_file:
            st.error("Please upload both your résumé and the job description.")
            st.stop()
        
        with st.spinner("Analyzing your documents and generating optimized résumé..."):
            resume_text = read_file(resume_file)
            job_text = read_file(job_file)
            
            if not resume_text or not job_text:
                st.error("Could not extract text from uploaded files. Please try again with different files.")
                st.stop()
            
            # Check for PDF extraction errors
            if resume_text.startswith("Error:") or job_text.startswith("Error:"):
                st.error(f"Error processing files: {resume_text if resume_text.startswith('Error:') else job_text}")
                st.stop()
            
            # Build prompt and call OpenAI
            prompt = build_prompt(resume_text, job_text, tone=tone)
            output = call_openai_chat(prompt, api_key_input)
            
            # Check for errors in API response
            if output.startswith("Error:"):
                st.error(output)
                st.stop()
            
            # Display success message
            st.success("Résumé successfully generated!")
            
            # Display generated resume
            st.subheader("📋 Generated Résumé Preview")
            st.text_area("", output, height=400, label_visibility="collapsed")
            
            # Create download buttons
            with tempfile.TemporaryDirectory() as tmpdir:
                if show_preview:
                    st.info("Review your résumé above before downloading.")
                
                col1, col2 = st.columns(2)
                
                with col1:
                    try:
                        docx_file = save_resume_docx(output, f"{tmpdir}/resume.docx")
                        with open(docx_file, "rb") as f:
                            st.download_button(
                                "📝 Download Word Document", 
                                f, 
                                file_name="tailored_resume.docx",
                                help="Download in Microsoft Word format for further editing",
                                use_container_width=True
                            )
                    except Exception as e:
                        st.error(f"Error creating Word document: {str(e)}")
                
                with col2:
                    try:
                        pdf_file = save_resume_pdf(output, f"{tmpdir}/resume.pdf")
                        with open(pdf_file, "rb") as f:
                            st.download_button(
                                "📄 Download PDF", 
                                f, 
                                file_name="tailored_resume.pdf",
                                help="Download in PDF format for easy sharing",
                                use_container_width=True
                            )
                    except Exception as e:
                        st.error(f"Error creating PDF document: {str(e)}")

if __name__ == "__main__":
    main()

# In[ ]:




